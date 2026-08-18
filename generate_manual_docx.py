"""
Generates HANA_Analysis_App_UserManual.docx in ECS Wiki style.
Run: python generate_manual_docx.py
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT = Path(r"C:\Users\I522148\OneDrive - SAP SE\SWAT\AI\Claude\working directory\Results\HANA_Analysis_App_UserManual.docx")

# ── Colours ───────────────────────────────────────────────────────────────────
C_DARK_BLUE   = RGBColor(0x00, 0x33, 0x66)
C_PRIMARY     = RGBColor(0x00, 0x57, 0xa8)
C_CYAN        = RGBColor(0x00, 0xb4, 0xe6)
C_WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
C_PAGE_BG     = RGBColor(0xe8, 0xf4, 0xfd)
C_LIGHT_BLUE  = RGBColor(0xf0, 0xf7, 0xff)
C_ROW_ALT     = RGBColor(0xf0, 0xf7, 0xff)
C_ROW_HOVER   = RGBColor(0xd0, 0xe8, 0xff)
C_GREEN_BG    = RGBColor(0xd4, 0xed, 0xda)
C_GREEN_FG    = RGBColor(0x15, 0x57, 0x24)
C_ORANGE_BG   = RGBColor(0xff, 0xf3, 0xcd)
C_ORANGE_FG   = RGBColor(0x85, 0x64, 0x04)
C_RED_BG      = RGBColor(0xfd, 0xe8, 0xe8)
C_RED_FG      = RGBColor(0xa0, 0x20, 0x20)
C_BLUE_BG     = RGBColor(0xcc, 0xe5, 0xff)
C_GREY_BG     = RGBColor(0xe9, 0xec, 0xef)
C_GREY_FG     = RGBColor(0x49, 0x50, 0x57)
C_PANEL_INFO  = RGBColor(0xe9, 0xf2, 0xff)
C_PANEL_TIP   = RGBColor(0xe3, 0xfc, 0xef)
C_PANEL_WARN  = RGBColor(0xff, 0xfa, 0xe6)
C_PANEL_ERR   = RGBColor(0xff, 0xeb, 0xe6)
C_BODY_TEXT   = RGBColor(0x17, 0x2b, 0x4d)
C_CODE_BG     = RGBColor(0xf4, 0xf5, 0xf7)
C_BORDER      = RGBColor(0xdf, 0xe1, 0xe6)


# ── XML helpers ───────────────────────────────────────────────────────────────
def hex_color(rgb: RGBColor) -> str:
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color(rgb))
    tcPr.append(shd)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def set_para_shading(para, rgb: RGBColor):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color(rgb))
    pPr.append(shd)


def add_para_border(para, side='left', color='0057a8', sz=24, space=360):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    el = OxmlElement(f'w:{side}')
    el.set(qn('w:val'), 'single')
    el.set(qn('w:sz'), str(sz))
    el.set(qn('w:space'), str(space))
    el.set(qn('w:color'), color)
    pBdr.append(el)
    pPr.append(pBdr)


def add_table_border(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'DFE1E6')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        cell = row.cells[col_idx]
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(int(width_cm * 567)))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)


# ── Style helpers ─────────────────────────────────────────────────────────────
def style_run(run, bold=False, italic=False, size=11, color=None, font='Segoe UI'):
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def add_heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    style_run(run, bold=True, size=14, color=C_DARK_BLUE)
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '0057a8')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    style_run(run, bold=True, size=12, color=C_PRIMARY)
    return p


def add_body(doc, text, size=10, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    style_run(run, size=size, color=color or C_BODY_TEXT)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.5)
    set_para_shading(p, RGBColor(0x1e, 0x1e, 0x2e))
    run = p.add_run(text)
    run.font.name  = 'Consolas'
    run.font.size  = Pt(9)
    run.font.color.rgb = RGBColor(0xcd, 0xd6, 0xf4)
    add_para_border(p, 'left', color='0057a8', sz=16, space=200)
    return p


def add_inline_code(para, text):
    run = para.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = C_DARK_BLUE
    run.font.highlight_color = None
    # add shading via rPr
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F4F5F7')
    rPr.append(shd)
    return run


def add_panel(doc, title, text, bg: RGBColor, border_color: str, title_color: RGBColor):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    set_para_shading(p, bg)
    add_para_border(p, 'left', color=border_color, sz=24, space=300)
    r1 = p.add_run(title + '  ')
    style_run(r1, bold=True, size=9, color=title_color)
    r2 = p.add_run(text)
    style_run(r2, size=9, color=C_BODY_TEXT)
    return p


def add_bullet(doc, text, level=0, num=False):
    p = doc.add_paragraph(style='List Bullet' if not num else 'List Number')
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    style_run(run, size=10, color=C_BODY_TEXT)
    return p


def header_row(table, cols):
    row = table.rows[0]
    for i, col in enumerate(cols):
        cell = row.cells[i]
        cell.text = ''
        set_cell_bg(cell, C_DARK_BLUE)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(col)
        style_run(run, bold=True, size=9, color=C_WHITE)


def data_row(table, row_idx, values, alt=False):
    row = table.rows[row_idx]
    bg = C_ROW_ALT if alt else C_WHITE
    for i, val in enumerate(values):
        cell = row.cells[i]
        set_cell_bg(cell, bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        # clear default text then add
        for p in cell.paragraphs:
            p.clear()
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        if isinstance(val, list):
            # list of (text, bold, color, code) tuples
            for item in val:
                txt, bold, color, is_code = item
                if is_code:
                    add_inline_code(p, txt)
                else:
                    r = p.add_run(txt)
                    style_run(r, bold=bold, size=9, color=color or C_BODY_TEXT)
        else:
            run = p.add_run(str(val))
            style_run(run, size=9, color=C_BODY_TEXT)


def add_step(doc, num, title, body):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    set_para_shading(p, C_LIGHT_BLUE)
    add_para_border(p, 'left', color='0057a8', sz=20, space=250)
    r1 = p.add_run(f"  {num}.  ")
    style_run(r1, bold=True, size=11, color=C_PRIMARY)
    r2 = p.add_run(f"{title} — ")
    style_run(r2, bold=True, size=10, color=C_DARK_BLUE)
    r3 = p.add_run(body)
    style_run(r3, size=10, color=C_BODY_TEXT)
    return p


# ── Banner ────────────────────────────────────────────────────────────────────
def add_banner(doc):
    # Full-width shaded paragraph as banner
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    set_para_shading(p, C_DARK_BLUE)
    r = p.add_run("SAP HANA Full Analysis Report Generator  |  User Manual  |  v2.0  |  RDE LAC")
    style_run(r, bold=True, size=13, color=C_WHITE)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(10)
    set_para_shading(p2, C_PRIMARY)
    r2 = p2.add_run("HECOPS  \u203a  RDE LAC  \u203a  Tools  \u203a  HANA Analysis App        "
                    "August 2026  |  Windows  |  Python 3.10+")
    style_run(r2, size=8, color=RGBColor(0xa8, 0xd0, 0xf0))
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT


# ─────────────────────────────────────────────────────────────────────────────
def build_doc():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    add_banner(doc)

    # ── 1. Overview ──────────────────────────────────────────────────────────
    add_heading1(doc, "1.  Overview")
    add_body(doc,
        "The SAP HANA Full Analysis Report Generator is a desktop tool that automates the creation "
        "of interactive HTML health check reports from HANA health check script output files. "
        "It submits the health check data to Claude via the local HAI proxy using a locked "
        "design-system prompt template, streams the response in real time, and opens the final "
        "report in Microsoft Edge.")

    add_panel(doc, "What it replaces:",
        "Manually copying health check output into a Claude chat, waiting for the response, "
        "saving the HTML and opening it — the app does all of this in one click.",
        C_PANEL_TIP, '5cb85c', C_GREEN_FG)

    add_heading2(doc, "Key Features")
    tbl = doc.add_table(rows=9, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    add_table_border(tbl)
    header_row(tbl, ["Feature", "Detail"])
    rows = [
        ("Streaming generation",  "Real-time token-by-token progress via Anthropic SDK — no timeouts"),
        ("Live progress bar",     "Percentage advances as tokens are generated (~40,000 token target)"),
        ("File validation",       "Size badge + warning for files exceeding 800 KB"),
        ("Auto SID detection",    "Extracts SID from file content; falls back to filename prefix"),
        ("Auto open in Edge",     "Report opens automatically when generation completes"),
        ("Cancel button",         "Stops the stream at any point"),
        ("HAI proxy auth",        "Reads credentials from ~/.claude/settings.json — no API key needed"),
        ("JS fallback",           "Injects collapse/expand JS if model truncates the script block"),
    ]
    for i, (feat, detail) in enumerate(rows, 1):
        data_row(tbl, i, [feat, detail], alt=(i % 2 == 0))

    # ── 2. Requirements ──────────────────────────────────────────────────────
    add_heading1(doc, "2.  Requirements")
    tbl2 = doc.add_table(rows=7, cols=3)
    tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT
    add_table_border(tbl2)
    header_row(tbl2, ["Component", "Requirement", "Status check"])
    req_rows = [
        ("Python",         "3.10 or later",             "python --version"),
        ("anthropic SDK",  "Latest (0.120+)",            "pip show anthropic"),
        ("HAI Proxy",      "Running at localhost:6655",  "Must be active before launching"),
        ("Claude Code CLI","Authenticated session",      "claude --version"),
        ("Microsoft Edge", "Any version",               "Installed at default path"),
        ("settings.json",  "Contains ANTHROPIC_AUTH_TOKEN + ANTHROPIC_BASE_URL",
                           "~/.claude/settings.json"),
    ]
    for i, (comp, req, check) in enumerate(req_rows, 1):
        data_row(tbl2, i, [comp, req,
                           [(check, False, C_BODY_TEXT, True)]], alt=(i % 2 == 0))

    add_heading2(doc, "Install the Anthropic SDK (first time only)")
    add_code(doc, "python -m pip install anthropic --user")

    # ── 3. Launching ─────────────────────────────────────────────────────────
    add_heading1(doc, "3.  Launching the App")
    add_body(doc, "Open a terminal and run:")
    add_code(doc,
        'cd "C:\\Users\\I522148\\OneDrive - SAP SE\\SWAT\\AI\\Claude\\working directory"\n'
        "python hana_analysis_app_v2.py")
    add_panel(doc, "Note:",
        "The HAI Proxy (localhost:6655) must be running before you generate a report. "
        "If the proxy is not running, generation will fail with an API connection error.",
        C_PANEL_INFO, '0057a8', C_PRIMARY)

    # ── 4. Quick Start ───────────────────────────────────────────────────────
    add_heading1(doc, "4.  Quick Start")
    steps = [
        ("Launch the app",
         "Run hana_analysis_app_v2.py from the working directory."),
        ("Select a health check file",
         "Click Browse… and navigate to HANA Health Check Reports\\. "
         "Select a .txt file. The size badge shows KB and estimated token count."),
        ("Review the file preview",
         "The first 60 lines are shown in the preview pane. "
         "Confirm it is a valid health check output before proceeding."),
        ("Click Generate Full Analysis Report",
         "The progress bar starts moving immediately as tokens stream in. "
         "Expected time: 3–8 minutes depending on file size."),
        ("Report opens automatically",
         "When generation completes, the HTML is saved to Results\\ "
         "and Microsoft Edge opens it automatically."),
    ]
    for i, (title, body) in enumerate(steps, 1):
        add_step(doc, i, title, body)

    # ── 5. Interface Guide ───────────────────────────────────────────────────
    add_heading1(doc, "5.  Interface Guide")

    add_heading2(doc, "Header Bar")
    tbl3 = doc.add_table(rows=4, cols=2)
    add_table_border(tbl3)
    header_row(tbl3, ["Element", "Description"])
    data_row(tbl3, 1, ["Title", "SAP HANA Full Analysis Report Generator"])
    data_row(tbl3, 2, ["Subtitle", "Shows: Powered by Anthropic SDK | v2.0 | model name"], alt=True)
    data_row(tbl3, 3, ["RDE LAC badge", "Team identifier — appears in header and footer"])

    add_heading2(doc, "Health Check Input File")
    tbl4 = doc.add_table(rows=4, cols=2)
    add_table_border(tbl4)
    header_row(tbl4, ["Element", "Description"])
    data_row(tbl4, 1, ["File name label", "Shows selected filename in dark blue when a file is loaded"])
    data_row(tbl4, 2, ["Size badge", "Blue = OK  |  Orange = Large (>600 KB)  |  Red = Blocked (>800 KB)"], alt=True)
    data_row(tbl4, 3, ["Browse… button", "Opens file picker defaulting to the Reports folder"])

    add_heading2(doc, "Generation Progress")
    add_body(doc,
        "The progress bar advances based on output tokens generated (~40,000 expected for a "
        "full report). The phase label below the bar shows the current activity and elapsed time.")

    add_heading2(doc, "Generation Log")
    add_body(doc, "Live status output example:")
    add_code(doc,
        "File    : HP4REY_HANA_Health_Check_report.txt\n"
        "Prompt  : 138,011 chars (~34,502 tokens)\n"
        "Model   : claude-sonnet-latest\n"
        "Output  : HP4REY_HANA_Analysis_Report_20260807_110656.html\n"
        "Streaming response...\n"
        "Stream complete in 4m 12s\n"
        "Saved: HP4REY_HANA_Analysis_Report_20260807_110656.html  (156 KB)\n\n"
        "\u2714  Report ready: ...Results\\HP4REY_HANA_Analysis_Report_20260807_110656.html")

    add_heading2(doc, "Action Buttons")
    tbl5 = doc.add_table(rows=4, cols=3)
    add_table_border(tbl5)
    header_row(tbl5, ["Button", "State", "Action"])
    data_row(tbl5, 1, ["\u25b6  Generate Full Analysis Report", "Enabled after file selection", "Starts generation"])
    data_row(tbl5, 2, ["\u25a0  Cancel", "Enabled during generation", "Stops the stream immediately"], alt=True)
    data_row(tbl5, 3, ["\U0001f310  Open Last Report in Edge", "Enabled after successful generation", "Re-opens the last saved report"])

    # ── 6. File Size Limits ──────────────────────────────────────────────────
    add_heading1(doc, "6.  File Size Limits")
    tbl6 = doc.add_table(rows=4, cols=4)
    add_table_border(tbl6)
    header_row(tbl6, ["File Size", "Tokens (approx)", "Status", "Behaviour"])
    data_row(tbl6, 1,
        [[("< 560 KB", True, C_GREEN_FG, False)],
         "< 140,000",
         [("OK", True, C_GREEN_FG, False)],
         "Blue badge — proceeds normally"])
    data_row(tbl6, 2,
        [[("560–800 KB", True, C_ORANGE_FG, False)],
         "140,000–200,000",
         [("Large", True, C_ORANGE_FG, False)],
         "Orange badge — may be slower, allowed"], alt=True)
    data_row(tbl6, 3,
        [[(" > 800 KB", True, C_RED_FG, False)],
         "> 200,000",
         [("Blocked", True, C_RED_FG, False)],
         "Warning dialog — Generate button disabled"])

    add_panel(doc, "Large files:",
        "Files above 800 KB typically contain raw SQL trace output or duplicated data. "
        "Extract only the relevant health check sections before processing.",
        C_PANEL_WARN, 'f0ad4e', RGBColor(0x85, 0x64, 0x04))

    # ── 7. Output Report ─────────────────────────────────────────────────────
    add_heading1(doc, "7.  Output Report")

    add_heading2(doc, "File Naming Convention")
    add_body(doc, "Reports are saved to:")
    add_code(doc, "Results\\{SID}_HANA_Analysis_Report_{YYYYMMDD_HHMMSS}.html")
    add_body(doc, "Example:  HP4REY_HANA_Analysis_Report_20260807_110656.html")

    add_heading2(doc, "Report Sections")
    tbl7 = doc.add_table(rows=12, cols=3)
    add_table_border(tbl7)
    header_row(tbl7, ["#", "Section", "Contents"])
    sections = [
        ("0",  "Environment Summary",          "KPI cards + system grid (SID, host, RAM, CPU, MDC, replication)"),
        ("1",  "HANA Database Version",        "Installed vs ECS standard, SPS match badge, upgrade recommendation"),
        ("2",  "ECS Standard Parameters",      "ERROR count, full ERROR table, notable deviations"),
        ("3",  "P1 Alerts – Last 40 Days",     "Alert count by DB, grouped by Alert ID, recommended actions"),
        ("4",  "Big Technical Tables",         "Count, memory+disk GB, bar chart, LOB flag, archiving notes"),
        ("5",  "Tables > 1.5 Billion Records", "Partition count; green box if none found"),
        ("6",  "Out-of-Memory Events",         "OOM count, memory utilisation grid, top allocators"),
        ("7",  "CPU Spikes \u2265 95%",         "Spike count, timestamp table with bar chart, avg CPU"),
        ("8",  "Failed Backups",               "Count, pattern analysis (same-day vs spread), full detail table"),
        ("9",  "Additional Risks",             "Auto-derived risk table with severity pills and actions"),
        ("10", "Minichecks",                   "CPU/Memory, expensive statements, largest tables, disk, I/O, savepoints"),
    ]
    for i, (num, title, contents) in enumerate(sections, 1):
        data_row(tbl7, i, [num, title, contents], alt=(i % 2 == 0))

    add_heading2(doc, "Interacting with the Report")
    items = [
        "Click any section header (or \u25b6 arrow) to expand / collapse it",
        "Click nav pills in the sticky top bar to jump to a section",
        "Hover over any underlined term (blue dashed) to see the glossary tooltip",
        "KPI cards at the top are clickable and scroll to their section",
    ]
    for item in items:
        add_bullet(doc, item)

    # ── 8. Troubleshooting ───────────────────────────────────────────────────
    add_heading1(doc, "8.  Troubleshooting")
    tbl8 = doc.add_table(rows=8, cols=3)
    add_table_border(tbl8)
    header_row(tbl8, ["Error / Symptom", "Cause", "Fix"])
    tr = [
        ("API error 401",
         "HAI proxy auth token expired or proxy not running",
         "Start the HAI proxy and re-authenticate your Claude Code session."),
        ("The operation timed out",
         "Non-streaming CLI call through proxy",
         "v2.0 uses SDK streaming — this should not occur. Restart the app."),
        ("Progress stuck at 10%",
         "Proxy not running or network issue",
         "Verify localhost:6655 is reachable. Restart proxy."),
        ("Sections don't expand on click",
         "Report truncated — script block missing",
         "v2.0 auto-injects toggleSection() JS. Regenerate the report."),
        ("File blocked (orange/red warning)",
         "File exceeds 800 KB",
         "Use a smaller health check file. See section 6."),
        ("name 'extract_sid' is not defined",
         "Corrupted app file",
         "Re-download hana_analysis_app_v2.py from the working directory."),
        ("Edge does not open",
         "Edge not at default path",
         "Update EDGE_PATH constant in the app file."),
    ]
    for i, (err, cause, fix) in enumerate(tr, 1):
        data_row(tbl8, i, [err, cause, fix], alt=(i % 2 == 0))

    # ── 9. FAQ ───────────────────────────────────────────────────────────────
    add_heading1(doc, "9.  FAQ")

    faqs = [
        ("Does this use my Anthropic API credits?",
         "No. The app routes through the local HAI proxy (localhost:6655) which uses your "
         "SAP corporate authentication. No personal API key or credits are required."),
        ("Where are reports saved?",
         "All reports are saved to: C:\\Users\\I522148\\OneDrive - SAP SE\\SWAT\\AI\\Claude\\working directory\\Results\\"),
        ("Can I process the same file twice?",
         "Yes. Each run generates a new timestamped file — previous reports are never overwritten."),
        ("Can I cancel mid-generation?",
         "Yes — click Cancel at any time. The stream stops immediately. No partial file is saved."),
        ("Which Claude model is used?",
         "The model is read from ANTHROPIC_MODEL in ~/.claude/settings.json. "
         "Currently: claude-sonnet-latest. It is shown in the app header."),
    ]
    for q, a in faqs:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(q)
        style_run(r, bold=True, size=10, color=C_DARK_BLUE)
        add_body(doc, a)

    add_heading2(doc, "How long does generation take?")
    tbl9 = doc.add_table(rows=4, cols=2)
    add_table_border(tbl9)
    header_row(tbl9, ["File Size", "Typical Time"])
    data_row(tbl9, 1, ["< 100 KB", "3–5 minutes"])
    data_row(tbl9, 2, ["100–400 KB", "5–10 minutes"], alt=True)
    data_row(tbl9, 3, ["400–800 KB", "10–20 minutes"])

    # ── 10. Changelog ────────────────────────────────────────────────────────
    add_heading1(doc, "10.  Changelog")
    tbl10 = doc.add_table(rows=3, cols=3)
    add_table_border(tbl10)
    header_row(tbl10, ["Version", "Date", "Changes"])
    data_row(tbl10, 1, [
        [("v2.0", True, C_PRIMARY, False)],
        "Aug 2026",
        "Switched to Anthropic SDK streaming (eliminates proxy timeouts)  \u00b7  "
        "Real-time progress bar  \u00b7  max_tokens raised to 64,000  \u00b7  "
        "File size limit raised to 800 KB  \u00b7  Auto JS injection fallback  \u00b7  "
        "Cancel button  \u00b7  Renamed to SAP HANA Full Analysis Report Generator  \u00b7  RDE LAC legend"
    ])
    data_row(tbl10, 2, [
        [("v1.0", True, C_GREY_FG, False)],
        "Aug 2026",
        "Initial release — claude CLI subprocess, indeterminate progress bar, API key input"
    ], alt=True)

    # ── Footer paragraph ─────────────────────────────────────────────────────
    doc.add_paragraph()
    p = doc.add_paragraph()
    set_para_shading(p, C_DARK_BLUE)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SAP HANA Full Analysis Report Generator  |  RDE LAC  |  User Manual v2.0  |  August 2026")
    style_run(r, size=8, color=RGBColor(0xa8, 0xd0, 0xf0))

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}  ({OUTPUT.stat().st_size // 1024} KB)")


if __name__ == "__main__":
    build_doc()
